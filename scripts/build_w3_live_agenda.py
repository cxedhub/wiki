#!/usr/bin/env python3
"""Build Workshop 3 Live Agenda docx (facilitator timing board).

Mirrors the structure of Workshop-2-Live-Agenda.docx but adapted to the
Workshop 3 agenda (https://www.cxedhub.com/pds/craft-pd-series/workshop-3-agenda.html)
and links each timeline row to the matching section heading in the shared
working Google Doc (1GaTY3haKOuy2qzUb-GfT4xE8yeRu9QpK).
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL


SHARED_DOC = (
    "https://docs.google.com/document/d/1GaTY3haKOuy2qzUb-GfT4xE8yeRu9QpK/"
    "edit?usp=drive_link&ouid=107503058538624505453&rtpof=true&sd=true"
)
PLANNED_AGENDA_HTML = "https://cxedhub.github.io/wiki/pds/craft-pd-series/workshop-3-agenda.html"
SLIDES = "https://cxedhub.github.io/wiki/pds/craft-pd-series/slides-w3.html"
ZOOM = "https://ucf.zoom.us/j/97106116299"
PRE_SURVEY = "https://ucf.qualtrics.com/jfe/form/SV_8G2kZYDDaRslsDY"
POST_SURVEY = "https://ucf.qualtrics.com/jfe/form/SV_5cnzJo6KUcwqa34"
MAKECODE = "https://makecode.microbit.org"
MICROPYTHON = "https://python.microbit.org"
IOT_TEMPLATE = "https://cxedhub.github.io/wiki/resources/iot-lesson-template.html"
SENSOR_REF = "https://cxedhub.github.io/wiki/pds/craft-pd-series/resources/sensor-reference-guide.html"
CTM_TEMPLATE = "https://cxedhub.github.io/wiki/pds/craft-pd-series/resources/ctm-template.html"
PROMPT_LIBRARY = "https://cxedhub.github.io/wiki/pds/craft-pd-series/resources/prompt-library.html"


def add_hyperlink(paragraph, url, text, bold=False):
    """Insert a hyperlink run into an existing paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    if bold:
        b = OxmlElement("w:b")
        rPr.append(b)

    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_col_widths(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for idx, w in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(w)


def write_cell(cell, text, bold=False, size=None):
    p = cell.paragraphs[0]
    if text is None:
        return p
    run = p.add_run(text)
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    return p


def cell_add_line(cell, text=None, link_url=None, link_text=None, bold=False, size=None):
    p = cell.add_paragraph()
    if text:
        r = p.add_run(text)
        if bold:
            r.bold = True
        if size:
            r.font.size = Pt(size)
    if link_url and link_text:
        add_hyperlink(p, link_url, link_text, bold=bold)
    return p


def build():
    doc = Document()

    # Tighter margins to keep the table readable on one page width
    section = doc.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # ── Header block ────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    t_run = title.add_run("Workshop 3 · LIVE AGENDA (Facilitator Mark-Up)")
    t_run.bold = True
    t_run.font.size = Pt(18)

    sub = doc.add_paragraph()
    s_run = sub.add_run(
        "Programming Edge/IoT Systems with AI · April 25, 2026 · "
        "8:30 AM – 12:00 PM PST · Virtual (Zoom)"
    )
    s_run.font.size = Pt(11)
    s_run.italic = True

    intro = doc.add_paragraph()
    intro.add_run(
        "Use during the live session to mark actual start/end times and in-flight "
        "adjustments. Each block links to the matching section in the Shared Group "
        "Workspace so facilitators can jump straight to where participants are working."
    ).font.size = Pt(10)

    # ── How to use ──────────────────────────────────────────────────────────
    doc.add_heading("🛠 How to use this document", level=2)
    for line in [
        "One row per agenda block. PLANNED times are pre-filled from the finalized agenda.",
        "As each block STARTS, type the actual clock time in the “ACTUAL start” cell.",
        "As each block ENDS, type the actual clock time in the “ACTUAL end” cell.",
        "The “Drift” cell is where you jot +5 / –3 / on-time — a quick signal for co-facilitator.",
        "The “Call” cell is your in-flight decision: HOLD · TRIM · SKIP · EXTEND.",
        "The “Notes” cell captures why — so post-session we can adjust future cohorts.",
        "Each Block link opens the Shared Workspace at the relevant section heading.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(line).font.size = Pt(10.5)

    # ── Quick Links ─────────────────────────────────────────────────────────
    doc.add_heading("🔗 Quick Links", level=2)
    row1 = doc.add_paragraph()
    add_hyperlink(row1, PLANNED_AGENDA_HTML, "Planned Agenda (HTML)")
    row1.add_run("   ·   ")
    add_hyperlink(row1, SLIDES, "Web Slides")
    row1.add_run("   ·   ")
    add_hyperlink(row1, SHARED_DOC, "Shared Workspace (Google Doc)")
    row1.add_run("   ·   ")
    add_hyperlink(row1, ZOOM, "Zoom Room")

    row2 = doc.add_paragraph()
    add_hyperlink(row2, PRE_SURVEY, "Pre-Survey")
    row2.add_run("   ·   ")
    add_hyperlink(row2, POST_SURVEY, "Post-Survey")
    row2.add_run("   ·   ")
    add_hyperlink(row2, MAKECODE, "MakeCode Simulator")
    row2.add_run("   ·   ")
    add_hyperlink(row2, MICROPYTHON, "MicroPython Simulator")

    row3 = doc.add_paragraph()
    add_hyperlink(row3, IOT_TEMPLATE, "IoT Lesson Template")
    row3.add_run("   ·   ")
    add_hyperlink(row3, SENSOR_REF, "Sensor Reference Guide")
    row3.add_run("   ·   ")
    add_hyperlink(row3, CTM_TEMPLATE, "CtM Template")
    row3.add_run("   ·   ")
    add_hyperlink(row3, PROMPT_LIBRARY, "Prompt Library")

    # ── Running Budget ──────────────────────────────────────────────────────
    doc.add_heading("⏱ Running Budget (target)", level=2)
    bp = doc.add_paragraph()
    bp.add_run(
        "Talk / Listen ≈ 30 min · You-Do ≈ 150 min · Breaks = 30 min · "
        "TOTAL = 210 min (3.5 h) · Talk:Do ≈ 1:5.0"
    ).font.size = Pt(10.5)

    # ── Agenda Timeline ─────────────────────────────────────────────────────
    doc.add_heading("📋 Agenda Timeline (mark up in real time)", level=2)

    headers = [
        "Plan start", "Plan end", "ACTUAL start", "ACTUAL end", "Dur",
        "Block (→ Shared Doc section)",
        "Key prompts / do’s",
        "CRAFT + Type",
        "Drift +/−",
        "Call · Notes",
    ]

    # Block label, shared-doc section link text (None = no section), description,
    # inline helper links [(label, url), ...]
    rows = [
        {
            "plan_start": "8:30", "plan_end": "8:45", "dur": "15",
            "block": "Opening: Simulator Check + Icebreaker + Pre-Survey",
            "section": None,
            "desc": "Open MakeCode, drop a show-string block, then drop a radio block and watch the second simulator appear. Chat drop: what have you built (or wished you could build) with a physical computing device? Pre-Survey link in chat.",
            "helpers": [("→ MakeCode Simulator", MAKECODE), ("→ Pre-Survey", PRE_SURVEY)],
            "craft": "Opening · Admin",
            "shade": "E8F0FE",  # blue
        },
        {
            "plan_start": "8:45", "plan_end": "9:00", "dur": "15",
            "block": "CRAFT Orientation + IoT as a Two-Device Conversation",
            "section": "🔍 Contextualize: Hardware Experience Check-In",
            "desc": "C→R→A→F→T intro (60s). Every IoT system is a conversation between sensor nodes + aggregator nodes. Shared Doc: name a real-world two-node scenario from a topic you teach. Star a favorite.",
            "helpers": [],
            "craft": "Contextualize · You Do",
            "shade": "E8F8F0",  # green-ish for Do
        },
        {
            "plan_start": "9:00", "plan_end": "9:15", "dur": "15",
            "block": "Poll + Physical Computing Lives in the Browser Too",
            "section": "🔄 Reframe: “Coding Is Physical, Not Just Screens”",
            "desc": "Poll (a/b/c/d): “To teach IoT I need…” Then reframe: MakeCode simulator IS a micro:bit; radio block spawns a second simulator; the .hex flashes to real hardware later unchanged. LLMs write the code; the skill is prompt + verify.",
            "helpers": [],
            "craft": "Reframe · Listen",
            "shade": "EAF4FB",  # listen blue
        },
        {
            "plan_start": "9:15", "plan_end": "9:30", "dur": "15",
            "block": "Breakout: Teaching Hardware Without Hardware",
            "section": "🔄 Reframe: “Coding Is Physical, Not Just Screens”",
            "desc": "In small groups: #1 barrier to IoT/physical computing in your building? Pick one simulator-first response and name the trade-off vs. hardware-first. Post in Shared Doc; share out.",
            "helpers": [],
            "craft": "Reframe · You Do",
            "shade": "E8F8F0",
        },
        {
            "plan_start": "9:30", "plan_end": "9:45", "dur": "15",
            "block": "Break #1",
            "section": None,
            "desc": "Stretch. Confirm MakeCode is loaded and at least one simulated micro:bit is visible. TAs handle chat questions.",
            "helpers": [],
            "craft": "Break",
            "shade": "F2F2F2",
        },
        {
            "plan_start": "9:45", "plan_end": "10:15", "dur": "30",
            "block": "Assemble: Server Room Guardian (live build + swap-in)",
            "section": "🛠️ Assemble Part A: Temperature Sensor — Collective Data Log",
            "desc": "Live build a temp sensor node + aggregator on radio group 7, LLM-generated, toggled Blocks ↔ JS. Then participants swap temperature for light/accel/sound/compass, build their own pair, save two .hex files.",
            "helpers": [("→ MakeCode Simulator", MAKECODE)],
            "craft": "Assemble · You Do",
            "shade": "E8F8F0",
        },
        {
            "plan_start": "10:15", "plan_end": "10:45", "dur": "30",
            "block": "Assemble: Enhancements + Brainstorm",
            "section": "💻 LLM-Assisted Code Gallery",
            "desc": "Watch two live upgrades: liveness ping (PING/ACK dead-node warning) + MakeCode Data panel graphing. Brainstorm new applications. Then build time: extend your pair (thresholds, alerts, bidirectional, aggregation). Log prompt → response → fix in Shared Doc.",
            "helpers": [],
            "craft": "Assemble · You Do",
            "shade": "E8F8F0",
        },
        {
            "plan_start": "10:45", "plan_end": "11:00", "dur": "15",
            "block": "Break #2 — Show & Tell",
            "section": "💻 LLM-Assisted Code Gallery",
            "desc": "Screenshot your paired simulators doing something interesting (sender left, aggregator right) — drop in chat or paste into the Code Gallery section. Brag a little.",
            "helpers": [],
            "craft": "Break",
            "shade": "F2F2F2",
        },
        {
            "plan_start": "11:00", "plan_end": "11:30", "dur": "30",
            "block": "Assemble: Level Up — Choose Your Track (A Mesh / B MicroPython)",
            "section": "🚀 Level-Up Track Log (Optional)",
            "desc": "Track A: add a 2nd sensor node in a 2nd browser tab; IDs per sender; aggregator tracks per-sender averages + drops. Track B: port Sensor + Aggregator to MicroPython (python.microbit.org), verify messages still flow.",
            "helpers": [("→ MicroPython Simulator (Track B)", MICROPYTHON)],
            "craft": "Assemble · You Do",
            "shade": "E8F8F0",
        },
        {
            "plan_start": "11:30", "plan_end": "11:45", "dur": "15",
            "block": "Fortify: Verify Simulated Data + Calibration Preview",
            "section": "✅ Fortify: Sensor Verification via CtM — Group Results",
            "desc": "(1) Drag temp slider to extremes — does avg respond? (2) Change radio group — aggregator goes silent? (3) Remove radio.setGroup — watch failure mode. Apply CtM: Task → Before → After → Takeaway. Preview the calibration surprise (CPU-die temp runs 3–8 °C hot).",
            "helpers": [("→ CtM Template", CTM_TEMPLATE)],
            "craft": "Fortify · You Do",
            "shade": "E8FAF6",  # teal
        },
        {
            "plan_start": "11:45", "plan_end": "12:00", "dur": "15",
            "block": "Transfer + Close: Two-Node IoT Lesson + Pair-Share + Post-Survey",
            "section": "🌐 Transfer: Shared IoT Lesson Idea Board",
            "desc": "CRAFT debrief — how W1 (AI co-pilot) + W2 (verification) + W3 (physical-in-a-browser) stack into an integrated STEM toolkit. Draft a two-node lesson in the IoT template. Pair-share: “first two-node lesson you’ll run.” Post-survey + lesson submitted = your micro:bit V2 kit ships this week.",
            "helpers": [("→ IoT Lesson Template", IOT_TEMPLATE), ("→ Post-Survey", POST_SURVEY)],
            "craft": "Transfer + Close · Admin",
            "shade": "E8F0FE",
        },
    ]

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        set_cell_bg(c, "1B2A4A")
        # white text
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Body rows
    for ri, row in enumerate(rows, start=1):
        r = table.rows[ri]

        write_cell(r.cells[0], row["plan_start"], size=9)
        write_cell(r.cells[1], row["plan_end"], size=9)
        write_cell(r.cells[2], "", size=9)  # ACTUAL start (blank)
        write_cell(r.cells[3], "", size=9)  # ACTUAL end (blank)
        write_cell(r.cells[4], row["dur"], size=9)

        # Block cell: title + optional section link
        block_cell = r.cells[5]
        bp = block_cell.paragraphs[0]
        br = bp.add_run(row["block"])
        br.bold = True
        br.font.size = Pt(9.5)
        if row["section"]:
            link_p = block_cell.add_paragraph()
            lr = link_p.add_run("→ Shared Doc: ")
            lr.font.size = Pt(8.5)
            add_hyperlink(link_p, SHARED_DOC, row["section"])

        # Key prompts / do’s
        desc_cell = r.cells[6]
        dp = desc_cell.paragraphs[0]
        dr = dp.add_run(row["desc"])
        dr.font.size = Pt(9)
        for label, url in row["helpers"]:
            hp = desc_cell.add_paragraph()
            add_hyperlink(hp, url, label)

        write_cell(r.cells[7], row["craft"], size=9)
        write_cell(r.cells[8], "", size=9)
        write_cell(r.cells[9], "", size=9)

        # Shade the row-type column (CRAFT + Type) to match agenda colors
        set_cell_bg(r.cells[7], row["shade"])

    # column widths (cm), totaling approx 18cm of usable page width
    set_col_widths(table, [1.3, 1.3, 1.5, 1.5, 0.9, 4.5, 5.0, 2.4, 1.2, 2.0])

    # ── In-flight decision key ──────────────────────────────────────────────
    doc.add_heading("🚦 In-flight decision key", level=2)
    doc.add_paragraph("If running ≥ 3 minutes LATE at any checkpoint, pick one:")
    for s in [
        "TRIM → compress the next Listen beat; core prompts stay, stories drop.",
        "SKIP → compress Break #2 (10:45) to 10 min, or cut the Track-A↔B cross-share at 11:25.",
        "EXTEND → cut 5 min from Level Up (11:00) so participants finish on their own.",
        "HOLD → push on schedule; no adjustment.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(s).font.size = Pt(10.5)

    doc.add_paragraph("If running ≥ 3 minutes EARLY:")
    for s in [
        "ADD chat share-outs in the current breakout (2–3 groups).",
        "SURFACE a Code Gallery entry (LLM win or LLM fail) live on shared screen.",
        "INVITE one participant to narrate a CtM takeaway on their sensor verification (90s).",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(s).font.size = Pt(10.5)

    # ── Post-Session Notes ──────────────────────────────────────────────────
    doc.add_heading("📝 Post-Session Notes", level=2)
    doc.add_paragraph(
        "What to pull forward into future cohorts / into the final report:"
    )
    notes_tbl = doc.add_table(rows=8, cols=3)
    notes_tbl.style = "Table Grid"
    notes_hdr = notes_tbl.rows[0]
    for i, h in enumerate(["Theme", "What worked", "What to change"]):
        cell = notes_hdr.cells[i]
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        set_cell_bg(cell, "1B2A4A")
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    themes = [
        "Pacing",
        "Engagement",
        "Simulator / MakeCode friction",
        "LLM prompting (Server Room Guardian)",
        "Level-Up track balance (Mesh vs. MicroPython)",
        "CtM adoption on sensor verification",
        "Other",
    ]
    for i, t in enumerate(themes, start=1):
        write_cell(notes_tbl.rows[i].cells[0], t, bold=True, size=10)

    # ── Co-Facilitator Backchannel ──────────────────────────────────────────
    doc.add_heading("💬 Co-Facilitator Backchannel", level=2)
    doc.add_paragraph(
        "Use this space for silent signals between facilitators during the "
        "session (e.g., “watch Group 3”, “chat flood”, “can you cover the next poll?”)."
    )
    bc = doc.add_table(rows=13, cols=4)
    bc.style = "Table Grid"
    bc_hdr = bc.rows[0]
    for i, h in enumerate(["Time", "From", "To", "Message"]):
        cell = bc_hdr.cells[i]
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        set_cell_bg(cell, "1B2A4A")
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    out = "/home/user/wiki/Workshop-3-Live-Agenda.docx"
    doc.save(out)
    print(f"wrote {out}")


if __name__ == "__main__":
    build()
